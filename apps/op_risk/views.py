from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from .models import (
    Macroprocess, Process, RiskCategory, ProbabilityLevel, ImpactLevel,
    Risk, RiskEvent, ActionPlan, Control, KeyRiskIndicator, OpRiskDocument,
    OperationalCapitalCalculation, Subprocess, Activity
)
from django.views.generic import CreateView, UpdateView
from django import forms
from django.http import Http404
from django.urls import reverse
from django.views.decorators.csrf import csrf_exempt

@login_required
def dashboard(request):
    macro_count = Macroprocess.objects.count()
    process_count = Process.objects.count()
    event_count = RiskEvent.objects.count()
    action_plan_count = ActionPlan.objects.filter(status__in=['OPEN', 'IN_PROGRESS', 'OVERDUE']).count()
    
    # Heatmap data (Residual Risk)
    heatmap_data = [[0 for _ in range(5)] for _ in range(5)]
    risks = Risk.objects.all().select_related('inherent_probability', 'inherent_impact', 'residual_probability', 'residual_impact')
    
    top_risks = []
    
    for risk in risks:
        prob = risk.residual_probability or risk.inherent_probability
        imp = risk.residual_impact or risk.inherent_impact
        
        if prob and imp:
            prob_idx = prob.level - 1
            imp_idx = imp.level - 1
            heatmap_data[prob_idx][imp_idx] += 1
            
            # Score: P * I
            score = prob.level * imp.level
            top_risks.append({
                'risk': risk,
                'score': score,
                'prob_level': prob.level,
                'imp_level': imp.level
            })
            
    # Sort top risks by score descending
    top_risks.sort(key=lambda x: x['score'], reverse=True)
    top_risks = top_risks[:15]

    # Risk by Category
    from django.db.models import Count
    categories = Risk.objects.values('category__name').annotate(total=Count('id')).order_by('-total')
    category_labels = [c['category__name'] or 'Sin Categoría' for c in categories]
    category_data = [c['total'] for c in categories]

    # Open Issues
    issues = ActionPlan.objects.values('status').annotate(total=Count('id'))
    issue_data = {
        'OPEN': 0, 'IN_PROGRESS': 0, 'COMPLETED': 0, 'OVERDUE': 0
    }
    for i in issues:
        issue_data[i['status']] = i['total']
    
    # Control Performance
    controls = Control.objects.all()
    control_data = {'High': 0, 'Medium': 0, 'Low': 0}
    for c in controls:
        avg_eff = (c.design_efficacy + c.operational_effectiveness) / 2
        if avg_eff >= 80:
            control_data['High'] += 1
        elif avg_eff >= 50:
            control_data['Medium'] += 1
        else:
            control_data['Low'] += 1

    from .models import OperationalCapitalCalculation
    capital_calc = OperationalCapitalCalculation.objects.order_by('-year').first()
    
    context = {
        'page_title': 'Panel de KPI de Gestión de Riesgos',
        'macro_count': macro_count,
        'process_count': process_count,
        'event_count': event_count,
        'action_plan_count': action_plan_count,
        'heatmap_data': heatmap_data,
        'top_risks': top_risks,
        'capital_calc': capital_calc,
        'category_labels': category_labels,
        'category_data': category_data,
        'issue_data': issue_data,
        'control_data': control_data,
    }
    return render(request, 'op_risk/dashboard.html', context)

@login_required
def process_list(request):
    macroprocesses = Macroprocess.objects.prefetch_related(
        'processes', 'owner_position', 'owner_area',
        'processes__subprocesses', 'processes__subprocesses__activities',
        'processes__owner_position', 'processes__owner_area'
    ).all()
    
    processes = Process.objects.select_related('macroprocess', 'owner_position', 'owner_area').all()
    subprocesses = Subprocess.objects.select_related('process', 'owner_position', 'owner_area').all()
    activities = Activity.objects.select_related('subprocess', 'owner_position', 'owner_area').all()
    
    context = {
        'page_title': 'Mapa Operativo (Procesos)',
        'macroprocesses': macroprocesses,
        'processes': processes,
        'subprocesses': subprocesses,
        'activities': activities,
    }
    return render(request, 'op_risk/process_list.html', context)

@login_required
def risk_matrix(request):
    risks = Risk.objects.select_related('process', 'category', 'owner', 'inherent_probability', 'inherent_impact', 'residual_probability', 'residual_impact').all()
    for r in risks:
        prob = r.residual_probability or r.inherent_probability
        imp = r.residual_impact or r.inherent_impact
        if prob and imp:
            r.score = prob.level * imp.level
        else:
            r.score = 0
            
    context = {'page_title': 'Matriz de Riesgos', 'risks': risks}
    return render(request, 'op_risk/generic_list.html', context)

@login_required
def control_matrix(request):
    from .models import Control
    controls = Control.objects.select_related('owner').all()
    context = {'page_title': 'Matriz de Controles', 'items': controls, 'type': 'control'}
    return render(request, 'op_risk/generic_list.html', context)

@login_required
def event_log(request):
    events = RiskEvent.objects.select_related('process', 'reported_by').all()
    context = {'page_title': 'Registro de Eventos', 'items': events, 'type': 'event'}
    return render(request, 'op_risk/generic_list.html', context)

@login_required
def kpi_kri(request):
    from .models import KeyRiskIndicator
    kris = KeyRiskIndicator.objects.select_related('process', 'risk', 'owner').all()
    context = {'page_title': 'KRIs y KPIs', 'items': kris, 'type': 'kri'}
    return render(request, 'op_risk/generic_list.html', context)

@login_required
def action_plans(request):
    plans = ActionPlan.objects.select_related('owner', 'risk', 'event').all()
    context = {'page_title': 'Planes de Acción', 'items': plans, 'type': 'plan'}
    return render(request, 'op_risk/generic_list.html', context)

@login_required
def documents(request):
    from .models import OpRiskDocument
    docs = OpRiskDocument.objects.select_related('uploaded_by', 'content_type').all()
    context = {'page_title': 'Gestión Documental', 'items': docs, 'type': 'doc'}
    return render(request, 'op_risk/generic_list.html', context)

@login_required
def capital_list(request):
    from .models import OperationalCapitalCalculation
    items = OperationalCapitalCalculation.objects.all().order_by('-year')
    context = {'page_title': 'Cálculos de Capital Operacional', 'items': items, 'type': 'capital'}
    return render(request, 'op_risk/generic_list.html', context)

@login_required
def reports(request):
    context = {'page_title': 'Reportes e Informes'}
    return render(request, 'op_risk/reports.html', context)

@login_required
def quick_report(request):
    from datetime import date
    from django.contrib import messages
    from django.shortcuts import redirect
    
    if request.method == 'POST':
        title = request.POST.get('title')
        event_type = request.POST.get('event_type')
        date_occurred = request.POST.get('date_occurred')
        process_id = request.POST.get('process_id')
        amount = request.POST.get('amount', 0)
        
        # Create Event
        RiskEvent.objects.create(
            title=title,
            event_type=event_type,
            date_occurred=date_occurred or date.today(),
            date_discovered=date.today(),
            process_id=process_id if process_id else None,
            amount=amount if amount else 0.0,
            reported_by=request.user,
            extra_data={'source': 'quick_report'}
        )
        messages.success(request, '¡Incidente reportado exitosamente! El equipo de riesgos lo evaluará.')
        return redirect('op_risk:dashboard')

    processes = Process.objects.all()
    event_types = RiskEvent.EVENT_TYPES
    return render(request, 'op_risk/quick_report.html', {
        'processes': processes,
        'event_types': event_types,
        'page_title': 'Reporte Rápido de Incidente'
    })

from django.http import JsonResponse
import json
import os
import google.generativeai as genai

@login_required
def ai_copilot(request):
    """
    Autocompletes the Risk form using Google Gemini based on the Process and Risk Name.
    """
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            entity_type = data.get('entity_type', 'riesgo')
            process_id = data.get('process_id')
            item_name = data.get('item_name', '').strip()
            
            # Retro-compatibility if JS still sends risk_name
            if not item_name and data.get('risk_name'):
                item_name = data.get('risk_name', '').strip()
            
            process_name = "Desconocido"
            if process_id:
                try:
                    process = Process.objects.get(id=process_id)
                    process_name = process.name
                except Process.DoesNotExist:
                    pass
            
            from catalogs.models import AIModulePrompt, SystemIntegration
            
            prompt_config = AIModulePrompt.objects.filter(module='op_risk', is_active=True).first()
            
            if prompt_config:
                provider_type = prompt_config.provider.provider
                api_key = prompt_config.provider.api_key
                model_name = prompt_config.model_name
                system_prompt_base = prompt_config.system_prompt
            else:
                integration = SystemIntegration.objects.filter(provider='gemini', is_active=True).first()
                api_key = integration.api_key if integration and integration.api_key else os.environ.get('GEMINI_API_KEY')
                provider_type = 'gemini'
                model_name = 'gemini-flash-latest'
                system_prompt_base = "Eres un experto en Riesgo Operacional (Basilea)."
            
            if not api_key:
                return JsonResponse({'status': 'error', 'message': 'Clave API no configurada para el proveedor de IA.'})
                
            categories = RiskCategory.objects.all()
            cat_list = ", ".join([f"ID {c.id}: {c.name}" for c in categories])
            probs = ProbabilityLevel.objects.all()
            prob_list = ", ".join([f"ID {p.id}: {p.name}" for p in probs])
            impacts = ImpactLevel.objects.all()
            impact_list = ", ".join([f"ID {i.id}: {i.name}" for i in impacts])

            prompt = f"{system_prompt_base}\n\nContexto:\n"
            
            if entity_type == 'riesgo':
                prompt += f"Tengo un proceso llamado '{process_name}'."
                if item_name:
                    prompt += f"\nHe identificado este riesgo: '{item_name}'. Completa los detalles de este riesgo específico.\n"
                else:
                    prompt += f"\nIdentifica un riesgo operacional común y significativo para este proceso.\n"
                    
                prompt += f"\nSelecciona los IDs numéricos más adecuados de estas opciones:\n"
                prompt += f"- Categorías: {cat_list}\n"
                prompt += f"- Probabilidades: {prob_list}\n"
                prompt += f"- Impactos: {impact_list}\n\n"
                
                prompt += (
                    "Genera un JSON estrictamente con la siguiente estructura (NO incluyas texto adicional ni bloques markdown ```json):\n"
                    "{\n"
                    '  "suggested_name": "Nombre corto y preciso del riesgo",\n'
                    '  "suggested_event": "Descripción del evento de riesgo (Qué podría salir mal y cómo)",\n'
                    '  "suggested_cause": "Causas raíz del riesgo",\n'
                    '  "suggested_consequence": "Consecuencias operativas, financieras o reputacionales",\n'
                    '  "suggested_category_id": ID numérico de la categoría,\n'
                    '  "suggested_prob_id": ID numérico de la probabilidad,\n'
                    '  "suggested_impact_id": ID numérico del impacto\n'
                    "}"
                )
            elif entity_type == 'control':
                prompt += f"Tengo un control llamado '{item_name}'."
                prompt += f"\nGenera los detalles para este control considerando las mejores prácticas de auditoría y riesgo operacional.\n"
                prompt += (
                    "Genera un JSON estrictamente con la siguiente estructura (NO incluyas texto adicional ni bloques markdown ```json):\n"
                    "{\n"
                    '  "suggested_name": "Nombre mejorado y formal del control",\n'
                    '  "suggested_description": "Descripción clara de cómo funciona el control",\n'
                    '  "suggested_type": "PREVENTIVE, DETECTIVE o CORRECTIVE",\n'
                    '  "suggested_periodicity": "Diaria, Semanal, Mensual, Trimestral, Anual o Continua",\n'
                    '  "suggested_design_efficacy": 85,\n'
                    '  "suggested_operative_efficacy": 85\n'
                    "}"
                )
            elif entity_type == 'evento':
                prompt += f"Tengo un evento de riesgo o incidente titulado '{item_name}'."
                prompt += f"\nGenera los detalles para este evento considerando la gestión de incidentes y riesgo operacional.\n"
                prompt += (
                    "Genera un JSON estrictamente con la siguiente estructura (NO incluyas texto adicional ni bloques markdown ```json):\n"
                    "{\n"
                    '  "suggested_title": "Título formal y descriptivo del evento",\n'
                    '  "suggested_event_type": "LOSS, NEAR_MISS, SYSTEM_FAILURE, o FRAUD",\n'
                    '  "suggested_root_cause": "Análisis profundo de la causa raíz que originó el evento",\n'
                    '  "suggested_immediate_action": "Acciones inmediatas recomendadas para detener el impacto y solucionar la falla"\n'
                    "}"
                )
            elif entity_type == 'kri':
                prompt += f"Tengo un Indicador Clave de Riesgo (KRI) llamado '{item_name}'."
                prompt += f"\nGenera los detalles para este KRI considerando las métricas de gestión de riesgo operacional.\n"
                prompt += (
                    "Genera un JSON estrictamente con la siguiente estructura (NO incluyas texto adicional ni bloques markdown ```json):\n"
                    "{\n"
                    '  "suggested_name": "Nombre mejorado y formal del KRI",\n'
                    '  "suggested_description": "Descripción clara de lo que mide este indicador y cómo se interpretaría su resultado",\n'
                    '  "suggested_green_threshold": 5,\n'
                    '  "suggested_yellow_threshold": 10,\n'
                    '  "suggested_red_threshold": 15\n'
                    "}"
                )
            elif entity_type == 'plan':
                prompt += f"Tengo un Plan de Acción o Hallazgo titulado '{item_name}'."
                prompt += f"\nGenera los detalles para este plan de acción considerando que busca mitigar un riesgo operacional o solucionar una deficiencia.\n"
                prompt += (
                    "Genera un JSON estrictamente con la siguiente estructura (NO incluyas texto adicional ni bloques markdown ```json):\n"
                    "{\n"
                    '  "suggested_title": "Título mejorado, claro y accionable",\n'
                    '  "suggested_description": "Descripción paso a paso de la acción correctiva o preventiva recomendada"\n'
                    "}"
                )
            elif entity_type == 'documento':
                prompt += f"Tengo un Documento de Riesgo cuyo borrador de nombre es '{item_name}'."
                prompt += f"\nGenera un nombre estandarizado y profesional para este documento, adecuado para una biblioteca de gestión de riesgo corporativo.\n"
                prompt += (
                    "Genera un JSON estrictamente con la siguiente estructura (NO incluyas texto adicional ni bloques markdown ```json):\n"
                    "{\n"
                    '  "suggested_title": "Nombre formal y estandarizado del documento"\n'
                    "}"
                )
            
            raw_text = ""
            
            if provider_type == 'gemini':
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel(model_name)
                response = model.generate_content(prompt)
                raw_text = response.text.strip()
            elif provider_type == 'openai':
                import requests
                headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
                payload = {
                    "model": model_name,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.5
                }
                res = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
                if res.status_code == 200:
                    raw_text = res.json()['choices'][0]['message']['content'].strip()
                else:
                    return JsonResponse({'status': 'error', 'message': f'Error OpenAI: {res.text}'}, status=400)
            elif provider_type == 'mistral':
                import requests
                headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
                payload = {
                    "model": model_name,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.5
                }
                res = requests.post("https://api.mistral.ai/v1/chat/completions", headers=headers, json=payload)
                if res.status_code == 200:
                    raw_text = res.json()['choices'][0]['message']['content'].strip()
                else:
                    return JsonResponse({'status': 'error', 'message': f'Error Mistral: {res.text}'}, status=400)
                
            if raw_text.startswith("```json"):
                raw_text = raw_text.strip("```json").strip("```").strip()
            elif raw_text.startswith("```"):
                raw_text = raw_text.strip("```").strip()
                
            suggestion = json.loads(raw_text)
            
            return JsonResponse({'status': 'success', 'data': suggestion})
        except Exception as e:
            error_msg = str(e)
            if '429' in error_msg or 'quota' in error_msg.lower():
                return JsonResponse({'status': 'error', 'message': '¡Uy! Nuestro asistente de Inteligencia Artificial está tomando un respiro rápido porque hay muchas consultas. Danos unos segunditos y vuelve a presionar el botón. ✨'}, status=400)
            return JsonResponse({'status': 'error', 'message': 'Algo salió mal al contactar la IA. Por favor, intenta nuevamente.'}, status=400)
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=405)

@login_required
def ai_chat_assistant(request):
    """
    AI Chat Assistant endpoint using Google Gemini.
    """
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            message = data.get('message', '')
            history = data.get('history', [])
            
            from catalogs.models import SystemIntegration
            integration = SystemIntegration.objects.filter(provider='gemini', is_active=True).first()
            api_key = integration.api_key if integration and integration.api_key else os.environ.get('GEMINI_API_KEY')
            
            if not api_key:
                return JsonResponse({'status': 'error', 'reply': 'Clave API de Gemini no configurada en el panel de Integraciones.'})
                
            genai.configure(api_key=api_key)
            
            system_instruction = (
                "Eres un asistente experto en Riesgo Operacional bajo los estándares de Basilea, trabajando para una COOPAC. "
                "Tu objetivo es ayudar al usuario a reportar un incidente llenando un formulario. "
                "Sé amable, conciso y profesional. En cada mensaje, ayúdalo a definir: 1) Tipo de evento (Fraude, Falla de Sistemas, Procesos, etc.), "
                "2) Proceso afectado, y 3) Monto estimado de pérdida. Ve paso a paso. No escribas respuestas muy largas."
            )
            
            model = genai.GenerativeModel('gemini-flash-latest', system_instruction=system_instruction)
            
            # Format history for Gemini API
            formatted_history = []
            for msg in history:
                if msg.get('role') in ['user', 'model']:
                    # We need to ensure we don't duplicate the last user message that was just sent
                    formatted_history.append({'role': msg['role'], 'parts': msg['parts']})
            
            # Start chat with all history except the latest message
            chat_history = formatted_history[:-1] if formatted_history else []
            chat = model.start_chat(history=chat_history)
            
            # Send the new message
            response = chat.send_message(message)
            
            return JsonResponse({'status': 'success', 'reply': response.text})
        except Exception as e:
            return JsonResponse({'status': 'error', 'reply': f'Error de IA: {str(e)}'})
            
    return JsonResponse({'status': 'error', 'reply': 'Método no permitido'}, status=405)

# --- Vistas Genéricas de Frontend ---

MODEL_MAPPING = {
    'macroproceso': Macroprocess,
    'proceso': Process,
    'subproceso': Subprocess,
    'actividad': Activity,
    'riesgo': Risk,
    'control': Control,
    'evento': RiskEvent,
    'kri': KeyRiskIndicator,
    'plan': ActionPlan,
    'documento': OpRiskDocument,
    'capital': OperationalCapitalCalculation,
}

URL_MAPPING = {
    'macroproceso': 'op_risk:process_list',
    'proceso': 'op_risk:process_list',
    'subproceso': 'op_risk:process_list',
    'actividad': 'op_risk:process_list',
    'riesgo': 'op_risk:risk_matrix',
    'control': 'op_risk:control_matrix',
    'evento': 'op_risk:event_log',
    'kri': 'op_risk:kpi_kri',
    'plan': 'op_risk:action_plans',
    'documento': 'op_risk:documents',
    'capital': 'op_risk:capital_list',
}

class BootstrapFormMixin:
    def get_form(self, form_class=None):
        form = super().get_form(form_class)
        for field_name, field in form.fields.items():
            if isinstance(field.widget, forms.CheckboxInput):
                field.widget.attrs['class'] = 'form-check-input'
            else:
                field.widget.attrs['class'] = 'form-control'
            # Añadir id específico para autocompletado con JS
            if getattr(self, 'model_name', None) == 'riesgo' and field_name == 'name':
                field.widget.attrs['id'] = 'riesgo-name-input'
        return form

class GenericCreateView(BootstrapFormMixin, CreateView):
    template_name = 'op_risk/generic_form.html'
    
    def dispatch(self, request, *args, **kwargs):
        self.model_name = self.kwargs.get('tipo')
        self.model = MODEL_MAPPING.get(self.model_name)
        if not self.model:
            raise Http404("Modelo no encontrado")
        self.fields = '__all__'
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['page_title'] = f"Crear {self.model._meta.verbose_name.title()}"
        context['model_name'] = self.model_name
        
        # Determine cancel URL
        cancel_url_name = URL_MAPPING.get(self.model_name, 'op_risk:dashboard')
        context['cancel_url'] = reverse(cancel_url_name)
        
        if self.model_name == 'riesgo':
            from catalogs.models import RiskCatalog
            context['risk_catalog'] = RiskCatalog.objects.values_list('name', flat=True)
            
        return context

    def get_success_url(self):
        success_url_name = URL_MAPPING.get(self.model_name, 'op_risk:dashboard')
        return reverse(success_url_name)

class GenericUpdateView(BootstrapFormMixin, UpdateView):
    template_name = 'op_risk/generic_form.html'
    
    def dispatch(self, request, *args, **kwargs):
        self.model_name = self.kwargs.get('tipo')
        self.model = MODEL_MAPPING.get(self.model_name)
        if not self.model:
            raise Http404("Modelo no encontrado")
        self.fields = '__all__'
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['page_title'] = f"Editar {self.model._meta.verbose_name.title()}"
        context['model_name'] = self.model_name
        
        # Determine cancel URL
        cancel_url_name = URL_MAPPING.get(self.model_name, 'op_risk:dashboard')
        context['cancel_url'] = reverse(cancel_url_name)
        
        if self.model_name == 'riesgo':
            from catalogs.models import RiskCatalog
            context['risk_catalog'] = RiskCatalog.objects.values_list('name', flat=True)
            
        return context

    def get_success_url(self):
        success_url_name = URL_MAPPING.get(self.model_name, 'op_risk:dashboard')
        return reverse(success_url_name)

from django.views.generic import DeleteView

class GenericDeleteView(DeleteView):
    template_name = 'op_risk/generic_confirm_delete.html'
    
    def dispatch(self, request, *args, **kwargs):
        self.model_name = self.kwargs.get('tipo')
        self.model = MODEL_MAPPING.get(self.model_name)
        if not self.model:
            raise Http404("Modelo no encontrado")
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['page_title'] = f"Eliminar {self.model._meta.verbose_name.title()}"
        context['model_name'] = self.model_name
        
        cancel_url_name = URL_MAPPING.get(self.model_name, 'op_risk:dashboard')
        context['cancel_url'] = reverse(cancel_url_name)
        
        return context

    def get_success_url(self):
        success_url_name = URL_MAPPING.get(self.model_name, 'op_risk:dashboard')
        return reverse(success_url_name)

@csrf_exempt
def global_ai_chat(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            user_message = data.get('message', '')
            current_url = data.get('url', '')
            history = data.get('history', [])
            
            api_key = os.environ.get('GEMINI_API_KEY')
            if not api_key:
                return JsonResponse({'status': 'error', 'reply': 'La clave de API GEMINI_API_KEY no está configurada en .env'})
            
            genai.configure(api_key=api_key)
            
            # Context builder
            context_instructions = ""
            if 'riesgo-operacional/procesos' in current_url:
                context_instructions = "El usuario está en el paso '1. Catálogo de Procesos'. Ayúdalo a definir o entender qué es un macroproceso, subproceso, responsable o criticidad según Basilea."
            elif 'riesgo-operacional/matriz-riesgos' in current_url:
                context_instructions = "El usuario está en el paso '2. Matriz de Riesgos'. Ayúdalo a identificar riesgos inherentes, probabilidad e impacto."
            elif 'riesgo-operacional/matriz-controles' in current_url:
                context_instructions = "El usuario está en el paso '3. Matriz de Controles'. Ayúdalo a establecer controles efectivos para mitigar el riesgo inherente."
            elif 'riesgo-operacional/registro-eventos' in current_url:
                context_instructions = "El usuario está en el paso '4. Registro de Eventos' o 'Incidentes'. Ayúdalo a documentar correctamente las pérdidas operacionales."
            elif 'riesgo-operacional/kpi-kri' in current_url:
                context_instructions = "El usuario está en el paso '5. KRIs / KPIs'. Ayúdalo a definir indicadores clave de riesgo y métricas de advertencia temprana."
            elif 'riesgo-operacional/planes-accion' in current_url:
                context_instructions = "El usuario está en el paso '6. Planes de Acción'. Ayúdalo a estructurar medidas correctivas y plazos."
            else:
                context_instructions = "El usuario está navegando por el Dashboard de Riesgo Operacional. Asístelo respondiendo a sus preguntas generales sobre el ciclo de gestión de riesgo."
                
            system_instruction = (
                "Eres un Asistente IA experto en Riesgo Operacional (metodología Basilea y ERM). "
                f"CONTEXTO DE NAVEGACIÓN ACTUAL: {context_instructions} "
                "Responde de forma clara, directa, educada y concisa (1 o 2 párrafos máximo). "
                "Si la pregunta no tiene relación con riesgo operacional o el sistema, indica amablemente cuál es tu propósito principal."
            )
            
            model = genai.GenerativeModel('gemini-flash-latest', system_instruction=system_instruction)
            
            formatted_history = []
            for msg in history:
                if msg['role'] in ['user', 'model']:
                    formatted_history.append({'role': msg['role'], 'parts': msg['parts']})
                    
            chat = model.start_chat(history=formatted_history)
            response = chat.send_message(user_message)
            
            return JsonResponse({
                'status': 'success',
                'reply': response.text
            })
            
        except Exception as e:
            return JsonResponse({'status': 'error', 'reply': f'Error de IA: {str(e)}'})
            
    return JsonResponse({'status': 'error', 'message': 'Método no permitido'}, status=405)

def get_executive_report_data():
    from django.db.models import Sum, Count, Q, F
    from datetime import date
    from .models import Risk, RiskEvent, Control, ActionPlan, KeyRiskIndicator, OperationalCapitalCalculation

    all_risks_approved = Risk.objects.filter(status__code='APPROVED').select_related('inherent_probability', 'inherent_impact', 'residual_probability', 'residual_impact')
    
    critical, high, medium, low = 0, 0, 0, 0
    top_risks_list = []
    
    for r in all_risks_approved:
        prob = r.residual_probability or r.inherent_probability
        imp = r.residual_impact or r.inherent_impact
        if prob and imp:
            score = prob.level * imp.level
            if score >= 15: critical += 1
            elif score >= 10: high += 1
            elif score >= 5: medium += 1
            else: low += 1
            
            top_risks_list.append({
                'risk': r,
                'score': score,
                'prob_level': prob.level,
                'imp_level': imp.level
            })

    top_risks_list.sort(key=lambda x: x['score'], reverse=True)
    top_risks_list = top_risks_list[:5]

    risk_profile = {
        'critical': critical,
        'high': high,
        'medium': medium,
        'low': low,
        'total': all_risks_approved.count()
    }

    current_year = date.today().year
    events_year = RiskEvent.objects.filter(date_occurred__year=current_year)
    total_losses = events_year.aggregate(total=Sum('amount'))['total'] or 0
    event_count = events_year.count()

    controls = Control.objects.all()
    control_stats = {
        'total': controls.count(),
        'effective': controls.filter(operational_effectiveness__gte=80).count(),
        'ineffective': controls.filter(operational_effectiveness__lt=80).count()
    }

    plans = ActionPlan.objects.all()
    plan_stats = {
        'total': plans.count(),
        'completed': plans.filter(status='COMPLETED').count(),
        'in_progress': plans.filter(status='IN_PROGRESS').count(),
        'overdue': plans.filter(status='OVERDUE').count()
    }

    kris = KeyRiskIndicator.objects.all()
    kri_stats = {'total': kris.count(), 'red': 0, 'yellow': 0, 'green': 0}
    red_kris_list = []
    for k in kris:
        latest = k.readings.first()
        if latest:
            val = latest.value
            if val > k.red_threshold:
                kri_stats['red'] += 1
                red_kris_list.append({'name': k.name, 'value': val, 'limit': k.red_threshold})
            elif val > k.green_threshold:
                kri_stats['yellow'] += 1
            else:
                kri_stats['green'] += 1
        else:
            kri_stats['green'] += 1

    capital_calc = OperationalCapitalCalculation.objects.first()

    # Detailed lists for all sections (as requested by user)
    all_processes = Process.objects.all().select_related('macroprocess')
    all_risks = Risk.objects.all().select_related('process', 'inherent_probability', 'inherent_impact', 'residual_probability', 'residual_impact')
    for r in all_risks:
        prob = r.residual_probability or r.inherent_probability
        imp = r.residual_impact or r.inherent_impact
        if prob and imp:
            r.score = prob.level * imp.level
        else:
            r.score = 0
            
    all_controls = Control.objects.all()
    all_events = RiskEvent.objects.all().order_by('-date_occurred')
    all_plans = ActionPlan.objects.all().order_by('commitment_date')
    all_kris = KeyRiskIndicator.objects.all()
    all_docs = OpRiskDocument.objects.all().order_by('-created_at')

    return {
        'report_date': date.today(),
        'risk_profile': risk_profile,
        'top_risks': top_risks_list,
        'total_losses': total_losses,
        'event_count': event_count,
        'control_stats': control_stats,
        'plan_stats': plan_stats,
        'kri_stats': kri_stats,
        'capital_calc': capital_calc,
        'all_processes': all_processes,
        'all_risks': all_risks,
        'all_controls': all_controls,
        'all_events': all_events,
        'all_plans': all_plans,
        'all_kris': all_kris,
        'all_docs': all_docs,
    }

@login_required
def executive_report(request):
    context = get_executive_report_data()
    context['page_title'] = 'Informe Ejecutivo de Gestión de Riesgo Operacional'
    return render(request, 'op_risk/executive_report.html', context)

from django.http import HttpResponse
from django.template.loader import render_to_string
from xhtml2pdf import pisa
import docx
from docx.shared import Pt, Inches

@login_required
def download_executive_pdf(request):
    context = get_executive_report_data()
    context['is_pdf'] = True # to modify some css
    html_string = render_to_string('op_risk/executive_report.html', context)
    
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="informe_ejecutivo_ro.pdf"'
    
    pisa_status = pisa.CreatePDF(html_string, dest=response)
    if pisa_status.err:
        return HttpResponse('Tuvimos errores generando el PDF <pre>' + html_string + '</pre>')
    return response

@login_required
def download_executive_docx(request):
    context = get_executive_report_data()
    doc = docx.Document()
    
    doc.add_heading('Informe Ejecutivo Gerencial', 0)
    doc.add_heading('Gestión de Riesgo Operacional (Basilea)', 1)
    doc.add_paragraph(f"Fecha de Emisión: {context['report_date'].strftime('%d/%m/%Y')}")
    
    doc.add_heading('1. Perfil de Riesgo Residual', level=2)
    p1 = doc.add_paragraph()
    p1.add_run(f"Críticos: {context['risk_profile']['critical']}\n")
    p1.add_run(f"Altos: {context['risk_profile']['high']}\n")
    p1.add_run(f"Medios: {context['risk_profile']['medium']}\n")
    p1.add_run(f"Bajos: {context['risk_profile']['low']}\n")
    p1.add_run(f"Total: {context['risk_profile']['total']}")
    
    doc.add_heading('2. Top 5 Riesgos (Severidad)', level=2)
    for tr in context['top_risks']:
        doc.add_paragraph(f"- {tr['risk'].name} | Proceso: {tr['risk'].process.name} | P: {tr['prob_level']} x I: {tr['imp_level']} | Score: {tr['score']}", style='List Bullet')

    doc.add_heading('3. Inventario de Procesos', level=2)
    for proc in context['all_processes']:
        doc.add_paragraph(f"- {proc.name} (Macroproceso: {proc.macroprocess.name if proc.macroprocess else 'N/A'})", style='List Bullet')

    doc.add_heading('4. Matriz de Riesgos', level=2)
    for risk in context['all_risks']:
        status_name = risk.status.name if risk.status else 'Borrador'
        doc.add_paragraph(f"- {risk.name} | Proceso: {risk.process.name} | Estado: {status_name} | Score Residual: {risk.score}", style='List Bullet')
        
    doc.add_heading('5. Matriz de Controles', level=2)
    for ctrl in context['all_controls']:
        doc.add_paragraph(f"- {ctrl.name} | Tipo: {ctrl.get_type_display()} | Efectividad: {ctrl.operational_effectiveness}%", style='List Bullet')
        
    doc.add_heading('6. Registro de Eventos', level=2)
    for event in context['all_events']:
        doc.add_paragraph(f"- {event.date_occurred.strftime('%d/%m/%Y')} | {event.title} | Pérdida: S/ {event.amount:.2f} | Tipo: {event.get_event_type_display()}", style='List Bullet')

    doc.add_heading('7. Indicadores Clave de Riesgo (KRI)', level=2)
    for kri in context['all_kris']:
        doc.add_paragraph(f"- {kri.name} | Límite Verde: {kri.green_threshold} | Límite Rojo: {kri.red_threshold}", style='List Bullet')

    doc.add_heading('8. Planes de Acción', level=2)
    for plan in context['all_plans']:
        doc.add_paragraph(f"- {plan.title} | Estado: {plan.get_status_display()} | Vencimiento: {plan.commitment_date.strftime('%d/%m/%Y')}", style='List Bullet')

    doc.add_heading('8. Gestión Documental', level=2)
    for d in context['all_docs']:
        doc.add_paragraph(f"- {d.title} | Versión: {d.version} | Fecha: {d.created_at.strftime('%d/%m/%Y')}", style='List Bullet')
         
    doc.add_heading('9. Matriz Integral de Riesgos y Controles', level=2)
    for risk in context['all_risks']:
        ctrls = risk.controls.all()
        if ctrls.exists():
            c_str = ", ".join([f"{c.name} ({c.operational_effectiveness}%)" for c in ctrls])
        else:
            c_str = "Sin controles mitigadores"
        doc.add_paragraph(f"- Proceso: {risk.process.name} | Riesgo: {risk.name} | Controles: {c_str} | Score: {risk.score}", style='List Bullet')

    doc.add_heading('10. Requerimiento de Patrimonio Efectivo', level=2)
    if context['capital_calc']:
        doc.add_paragraph(f"S/ {context['capital_calc'].calculated_capital:.2f}")
    else:
        doc.add_paragraph("No calculado")
        
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="informe_ejecutivo_ro.docx"'
    doc.save(response)
    return response
